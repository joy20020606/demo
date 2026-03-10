from docx import Document
from docx.shared import Pt

def create_demo():
    doc = Document()
    doc.add_heading('測試翻譯文件 (Test Doc)', 0)

    p = doc.add_paragraph('這是一個測試段落，包含 ')
    p.add_run('粗體文字').bold = True
    p.add_run(' 和 ')
    p.add_run('斜體文字').italic = True
    p.add_run('。')
   
    doc.add_paragraph('第二個段落：我們要把中文翻譯成其他語言，同時保留格式。')

    table = doc.add_table(rows=2, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '項目 (Item)'
    hdr_cells[1].text = '內容 (Content)'
    
    row_cells = table.rows[1].cells
    row_cells[0].text = '手機'
    row_cells[1].text = '旗艦款智能手機'

    doc.save('demo_input.docx')
    print("已建立測試文件: demo_input.docx")

if __name__ == "__main__":
    create_demo()
