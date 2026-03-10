import re
import shutil
import os
from docx import Document
from deep_translator import GoogleTranslator

try:
    from opencc import OpenCC
except ImportError:
    OpenCC = None

class DocxTranslator:
    def __init__(self, target_lang='ms'):
        self.target_lang = target_lang
        self.translator_auto = GoogleTranslator(source='auto', target=target_lang)
        self.translator_zh_tw = GoogleTranslator(source='zh-TW', target=target_lang)
        self.translator_zh_cn = GoogleTranslator(source='zh-CN', target=target_lang)
        self.t2s_converter = OpenCC('t2s') if OpenCC else None
        # 擴充中文範圍：涵蓋 CJK Unified Ideographs、Extension A、相容區
        # 這可避免像「體」這種 > U+9FA5 的繁中常用字被漏掉。
        self.chinese_re = re.compile(r'[\u3400-\u4DBF\u4E00-\u9FFF\uF900-\uFAFF]+')
        self.cache = {}  # 加入快取，避免重複翻譯相同的字詞

    def _translate_with_fallback(self, chinese_text):
        """優先直接翻譯，失敗或回傳原文時再走繁轉簡 fallback。"""
        attempts = []
        simplified = self._to_simplified(chinese_text)

        # 先嘗試原文，再嘗試繁轉簡後文本。
        attempts.append((self.translator_auto, chinese_text))
        attempts.append((self.translator_zh_tw, chinese_text))
        if simplified != chinese_text:
            attempts.append((self.translator_zh_cn, simplified))
            attempts.append((self.translator_auto, simplified))

        for translator, src_text in attempts:
            try:
                translated = translator.translate(src_text)
                # deep-translator 偶爾會把 Google 錯誤頁文字當成翻譯結果，需忽略。
                if translated and self._is_valid_translation(chinese_text, translated):
                    return translated
            except Exception:
                continue

        return chinese_text

    def _to_simplified(self, text):
        """Convert Traditional Chinese to Simplified when OpenCC is available."""
        if self.t2s_converter is None:
            return text
        try:
            return self.t2s_converter.convert(text)
        except Exception:
            return text

    def _is_valid_translation(self, original_text, translated_text):
        if not translated_text or not translated_text.strip():
            return False

        translated = translated_text.strip()
        original = original_text.strip()

        # 無效回應關鍵字（Google 錯誤頁常見內容）
        invalid_markers = [
            "Error 500",
            "Server Error",
            "That's an error",
            "That’s an error",
            "Please try again later",
        ]
        if any(marker in translated for marker in invalid_markers):
            return False

        return translated != original

    def contains_chinese(self, text):
        if not text:
            return False
        return bool(self.chinese_re.search(text))

    def translate_mixed_text(self, text):
        if not text:
            return text
            
        def translate_match(match):
            chinese_text = match.group(0)
            
            # 從快取讀取，加速翻譯
            if chinese_text in self.cache:
                return self.cache[chinese_text]
                
            try:
                translated = self._translate_with_fallback(chinese_text)
                if translated:
                    # 只快取成功翻譯結果，避免把失敗時的原文鎖死在快取中。
                    if translated.strip() != chinese_text.strip():
                        self.cache[chinese_text] = translated
                    return translated
                return chinese_text
            except Exception as e:
                print(f"翻譯失敗: {chinese_text}, 錯誤: {e}")
                return chinese_text
                
        # 只把連續中文的部分抓出來翻譯並替換
        return self.chinese_re.sub(translate_match, text)

    def process_paragraph(self, paragraph):
        """處理單個段落的替換邏輯"""
        if not paragraph or not paragraph.text.strip():
            return

        # 核心檢查：如果這個段落完全沒有中文，我們就完全不去碰它
        if not self.contains_chinese(paragraph.text):
            return

        # 關鍵修改：針對每個「跑段(Run)」個別去判斷並取代中文
        # 這樣做可以 100% 完美保留每一個跑段的格式（包含字體顏色、大小、粗體等），且不動到英文
        for run in paragraph.runs:
            if self.contains_chinese(run.text):
                run.text = self.translate_mixed_text(run.text)

        # 補強：有些文字不會出現在 paragraph.runs（例如特定嵌套節點），
        # 直接掃描段落內所有 w:t 節點，避免漏翻。
        for text_node in paragraph._p.xpath(".//*[local-name()='t']"):
            if self.contains_chinese(text_node.text):
                text_node.text = self.translate_mixed_text(text_node.text)

    def process_xml_text_nodes(self, root_element):
        """直接處理 XML 文字節點，涵蓋常規 API 可能漏掉的文字。"""
        for text_node in root_element.xpath(".//*[local-name()='t']"):
            if self.contains_chinese(text_node.text):
                text_node.text = self.translate_mixed_text(text_node.text)

    def process_table(self, table):
        """遞迴處理表格與巢狀表格。"""
        for row in table.rows:
            if not row:
                continue
            for cell in row.cells:
                for para in cell.paragraphs:
                    self.process_paragraph(para)
                # 處理巢狀表格，避免漏翻。
                for nested_table in cell.tables:
                    self.process_table(nested_table)

    def translate_docx(self, input_path, output_path=None):
        if not os.path.exists(input_path):
            print(f"找不到檔案: {input_path}")
            return

        if output_path is None:
            name, ext = os.path.splitext(input_path)
            output_path = f"{name}_{self.target_lang}{ext}"

        # 物理複製一份
        shutil.copy2(input_path, output_path)
        doc = Document(output_path)

        # A. 翻譯主本文
        print("正在處理本文段落...")
        for para in doc.paragraphs:
            self.process_paragraph(para)
        self.process_xml_text_nodes(doc.element.body)

        # B. 翻譯表格
        print("正在處理表格...")
        for table in doc.tables:
            self.process_table(table)

        # C. 翻譯頁首與頁尾 (Headers & Footers)
        print("正在處理頁首頁尾...")
        try:
            for section in doc.sections:
                for header_para in section.header.paragraphs:
                    self.process_paragraph(header_para)
                for header_table in section.header.tables:
                    self.process_table(header_table)
                for footer_para in section.footer.paragraphs:
                    self.process_paragraph(footer_para)
                for footer_table in section.footer.tables:
                    self.process_table(footer_table)

                # 再用 XML 節點層級補齊頁首/頁尾可能漏掉的文字節點。
                self.process_xml_text_nodes(section.header._element)
                self.process_xml_text_nodes(section.footer._element)
        except Exception:
            pass

        doc.save(output_path)
        print(f"\n翻譯完成！產出檔案: {output_path}")
        return output_path
    
